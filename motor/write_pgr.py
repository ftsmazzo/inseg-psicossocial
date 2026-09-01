from __future__ import annotations

import copy
import logging
import re
import unicodedata
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table, _Row

from motor.llm import is_robotic_danos
from motor.models import ProposedLine
from motor.pgr_cronograma import apply_psicossocial_cronogram
from motor.pgr_narrative import apply_psicossocial_narratives

logger = logging.getLogger(__name__)

_SAFE_DANOS_FALLBACK = "Agravos Ocupacionais SST a Definir — Revisar Manualmente"
_YELLOW_FILL = "FFFF00"


def _set_cell_text(cell, text: str) -> None:
    """Replace cell text preserving first paragraph style when possible."""
    text = text or ""
    if cell.paragraphs:
        # clear all paras except first
        p0 = cell.paragraphs[0]
        if p0.runs:
            p0.runs[0].text = text
            for run in p0.runs[1:]:
                run.text = ""
        else:
            p0.text = text
        # remove extra paragraphs
        for p in cell.paragraphs[1:]:
            p._element.getparent().remove(p._element)
    else:
        cell.text = text


def _clone_row(table: Table, row_index: int) -> _Row:
    row = table.rows[row_index]
    new_tr = copy.deepcopy(row._tr)
    row._tr.addnext(new_tr)
    return table.rows[row_index + 1]


def apply_lines_to_pgr(
    source_docx: str | Path,
    output_docx: str | Path,
    lines: list[ProposedLine],
    *,
    only_accepted: bool = True,
) -> dict:
    """Update/insert psychosocial APRHO rows. Returns stats."""
    source_docx = Path(source_docx)
    output_docx = Path(output_docx)
    doc = Document(str(source_docx))

    applied = 0
    skipped = 0
    # group by table
    by_table: dict[int, list[ProposedLine]] = {}
    for ln in lines:
        if only_accepted and ln.status.value not in {"Definitivo", "Preliminar", "Proposta"}:
            skipped += 1
            continue
        # Still write Preliminar/Proposta — technician may filter in UI later
        by_table.setdefault(ln.aprho_table_index, []).append(ln)

    for t_index, group in by_table.items():
        if t_index < 0 or t_index >= len(doc.tables):
            skipped += len(group)
            continue
        table = doc.tables[t_index]
        # process update first, then inserts
        updates = [g for g in group if g.action == "update_existing"]
        inserts = [g for g in group if g.action != "update_existing"]

        for ln in updates:
            row_i = ln.psico_row_index
            if row_i is None or row_i >= len(table.rows):
                inserts.append(ln)
                continue
            _write_row(table.rows[row_i], ln)
            applied += 1

        # For inserts: clone psico row or last data row
        for ln in inserts:
            anchor = ln.psico_row_index
            if anchor is None:
                anchor = _find_insert_anchor(table)
            elif anchor >= len(table.rows):
                anchor = _find_insert_anchor(table)
            else:
                cat = _row_category(table.rows[anchor])
                if _is_acidente_or_mecanico(cat):
                    anchor = _find_insert_anchor(table)
            new_row = _clone_row(table, anchor)
            _set_row_fill(new_row, _YELLOW_FILL)
            _write_row(new_row, ln)
            applied += 1

    narrative = apply_psicossocial_narratives(doc)
    cronogram = apply_psicossocial_cronogram(doc, lines)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))
    return {
        "applied": applied,
        "skipped": skipped,
        "output": str(output_docx),
        "narrative": narrative,
        "cronogram": cronogram,
    }


def _norm_cat(text: str) -> str:
    t = unicodedata.normalize("NFKC", text or "")
    return re.sub(r"\s+", " ", t.replace("\xa0", " ")).strip().lower()


def _row_category(row) -> str:
    if not row.cells:
        return ""
    return (row.cells[0].text or "").strip()


def _is_psico_category(cat: str) -> bool:
    return "psicos" in _norm_cat(cat)


def _is_ergonomico_category(cat: str) -> bool:
    n = _norm_cat(cat)
    return "ergon" in n and "psicos" not in n


def _is_acidente_or_mecanico(cat: str) -> bool:
    n = _norm_cat(cat)
    return any(k in n for k in ("mecan", "mecân", "acident", "biolog"))


def _set_row_fill(row, fill_hex: str) -> None:
    for cell in row.cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), fill_hex)


def _find_insert_anchor(table: Table) -> int:
    """
    Linha após a qual inserir psicossocial — sempre abaixo de Ergonômico/psico (amarelo),
    nunca após Acidentes/Mecânico (azul).
    """
    psico_idx: int | None = None
    last_ergo_idx: int | None = None
    first_acidente_idx: int | None = None

    for i, row in enumerate(table.rows):
        cat = _row_category(row)
        if not cat or _norm_cat(cat) in {"categoria", "perigos"}:
            continue
        if _is_psico_category(cat):
            psico_idx = i
        elif _is_ergonomico_category(cat):
            last_ergo_idx = i
        elif _is_acidente_or_mecanico(cat) and first_acidente_idx is None:
            first_acidente_idx = i

    if psico_idx is not None:
        return psico_idx
    if last_ergo_idx is not None:
        return last_ergo_idx
    if first_acidente_idx is not None and first_acidente_idx > 0:
        return first_acidente_idx - 1
    return max(len(table.rows) - 2, 1)


def _find_psico_index(table: Table) -> int | None:
    idx = _find_insert_anchor(table)
    cat = _row_category(table.rows[idx]) if idx < len(table.rows) else ""
    if _is_psico_category(cat) or _is_ergonomico_category(cat):
        return idx
    return None


def _write_row(row, ln: ProposedLine) -> None:
    danos = ln.danos or ""
    if is_robotic_danos(danos):
        logger.warning(
            "DOCX safety-net: danos clínico/robótico bloqueado GHE=%s danos=%r",
            ln.ghe_numero,
            danos,
        )
        danos = _SAFE_DANOS_FALLBACK
    values = [
        ln.categoria,
        ln.agente,
        ln.exposicao,
        ln.causa_fonte,
        ln.trajetoria,
        danos,
        str(ln.grau_exposicao),
        str(ln.grau_efeito),
        ln.potencial,
        ln.controles,
    ]
    cells = row.cells
    # table has 10 logical cols but merges may duplicate; write first 10 unique positions
    # python-docx returns merged cells repeatedly — write by index 0..9 carefully
    written = set()
    for idx, val in enumerate(values):
        if idx >= len(cells):
            break
        # skip if this cell object already written (merge)
        cell_id = id(cells[idx]._tc)
        if cell_id in written:
            continue
        _set_cell_text(cells[idx], val)
        written.add(cell_id)
