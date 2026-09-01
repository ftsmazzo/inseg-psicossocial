"""Teste de integração das narrativas em fixture sintética."""

from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[2]))

from docx.oxml.ns import qn

from motor.pgr_narrative import apply_psicossocial_narratives
from motor.tests.fixtures.build_minimal_narrative_doc import build_minimal_narrative_doc


def _find(doc, needle: str):
    for p in doc.paragraphs:
        if needle.lower() in p.text.lower():
            return p
    return None


def _run_fonts(p) -> set[str]:
    out: set[str] = set()
    for run in p.runs:
        if not run.text.strip():
            continue
        r_pr = run._element.find(qn("w:rPr"))
        if r_pr is None:
            out.add("inherit")
            continue
        rf = r_pr.find(qn("w:rFonts"))
        out.add("inherit" if rf is None else (rf.get(qn("w:ascii")) or "inherit"))
    return out


def _space_after(p) -> int | None:
    if p._p.pPr is None:
        return None
    sp = p._p.pPr.find(qn("w:spacing"))
    if sp is None:
        return None
    val = sp.get(qn("w:after"))
    return int(val) if val else None


def test_apply_all_narratives_on_minimal_fixture():
    doc = build_minimal_narrative_doc()
    count_before = len(doc.paragraphs)
    sect_before = sum(
        1 for p in doc.paragraphs if p._p.pPr is not None and p._p.pPr.sectPr is not None
    )

    results = apply_psicossocial_narratives(doc)
    assert all(v == "applied" for v in results.values())
    assert len(doc.paragraphs) == count_before + 10  # 1+1+1+6+2 inserts; 5.2 append
    assert sum(
        1 for p in doc.paragraphs if p._p.pPr is not None and p._p.pPr.sectPr is not None
    ) == sect_before

    p55 = _find(doc, "Para os fatores psicossociais")
    assert p55 is not None and p55.style.name == "Body Text"
    assert "Cambria" not in _run_fonts(p55)
    assert _space_after(p55) in (None, 635)

    p53 = _find(doc, "O acompanhamento da saúde ocupacional")
    assert p53 is not None and p53.style.name == "List Paragraph"
    assert "Cambria" not in _run_fonts(p53)

    intro = _find(doc, "Em casos de crise emocional")
    assert intro is not None
    assert intro._p.pPr is None or intro._p.pPr.numPr is None

    bullet = _find(doc, "Manter a calma e afastar")
    assert bullet is not None and bullet._p.pPr is not None and bullet._p.pPr.numPr is not None

    ref = _find(doc, "Relatório de Avaliação Psicossocial")
    assert ref is not None and "Cambria" not in _run_fonts(ref)

    results2 = apply_psicossocial_narratives(doc)
    assert all(v == "skipped" for v in results2.values())


if __name__ == "__main__":
    test_apply_all_narratives_on_minimal_fixture()
    print("ok")
