"""Testes unitários de insert/append sanitizado (sem depender de modelos externos)."""

from __future__ import annotations

import tempfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from motor.pgr_docx_utils import (
    append_to_paragraph,
    insert_paragraph_after,
    paragraph_has_numpr,
)


def _add_numpr(p: Paragraph, num_id: int = 1, ilvl: int = 0) -> None:
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl_el)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)


def _add_sect_pr(p: Paragraph) -> None:
    p_pr = p._p.get_or_add_pPr()
    p_pr.append(OxmlElement("w:sectPr"))


def _add_spacing(p: Paragraph, after: int) -> None:
    p_pr = p._p.get_or_add_pPr()
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:after"), str(after))
    p_pr.append(sp)


def _add_cambria_run(p: Paragraph, text: str) -> None:
    run = p.add_run(text)
    r_pr = run._element.get_or_add_rPr()
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Cambria")
    fonts.set(qn("w:hAnsi"), "Cambria")
    r_pr.append(fonts)


def _run_fonts(p: Paragraph) -> set[str]:
    out: set[str] = set()
    for run in p.runs:
        r_pr = run._element.find(qn("w:rPr"))
        if r_pr is None:
            out.add("inherit")
            continue
        rf = r_pr.find(qn("w:rFonts"))
        out.add("inherit" if rf is None else (rf.get(qn("w:ascii")) or "inherit"))
    return out


def test_append_preserves_single_paragraph():
    doc = Document()
    p = doc.add_paragraph("Texto base.", style="Body Text")
    append_to_paragraph(p, " Sufixo extra.")
    assert len(doc.paragraphs) == 1
    assert "Sufixo extra" in p.text
    assert p.text.startswith("Texto base")


def _space_after_twips(p: Paragraph) -> int | None:
    if p._p.pPr is None:
        return None
    sp = p._p.pPr.find(qn("w:spacing"))
    if sp is None:
        return None
    val = sp.get(qn("w:after"))
    return int(val) if val else None


def test_insert_strips_sectpr_and_inflated_spacing():
    doc = Document()
    anchor = doc.add_paragraph("Item lista", style="List Paragraph")
    _add_numpr(anchor)
    _add_spacing(anchor, 152400)
    _add_cambria_run(anchor, "Item lista")

    tpl = doc.add_paragraph("Template inflado", style="Heading 1")
    _add_sect_pr(tpl)
    _add_spacing(tpl, 200000)
    _add_numpr(tpl)

    new_p = insert_paragraph_after(
        tpl,
        "Inserido psico",
        anchor,
        keep_numpr=True,
        spacing_from=anchor,
    )

    assert new_p._p.pPr is None or new_p._p.pPr.sectPr is None
    sa = _space_after_twips(new_p)
    assert sa is None or sa <= 152400
    assert paragraph_has_numpr(new_p)
    assert "Cambria" not in _run_fonts(new_p)


def test_insert_intro_without_numpr():
    doc = Document()
    after = doc.add_paragraph("Último bullet elétrico", style="Heading 1")
    _add_numpr(after)
    tpl = doc.add_paragraph("Template bullet", style="Heading 1")
    _add_numpr(tpl)

    intro = insert_paragraph_after(
        tpl,
        "Intro psicossocial:",
        after,
        keep_numpr=False,
        spacing_from=after,
    )
    assert not paragraph_has_numpr(intro)


def test_roundtrip_save():
    doc = Document()
    anchor = doc.add_paragraph("Corpo descritivo Calibri.", style="Body Text")
    tpl = doc.add_paragraph("Template", style="Body Text")
    _add_cambria_run(tpl, "x")
    insert_paragraph_after(tpl, "Parágrafo inserido.", anchor)

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "t.docx"
        doc.save(str(path))
        doc2 = Document(str(path))
        assert any("Parágrafo inserido" in p.text for p in doc2.paragraphs)


if __name__ == "__main__":
    test_append_preserves_single_paragraph()
    test_insert_strips_sectpr_and_inflated_spacing()
    test_insert_intro_without_numpr()
    test_roundtrip_save()
    print("ok")
