"""Testes de posição e cor das linhas APRHO psicossociais."""

from __future__ import annotations

import tempfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from motor.models import LineStatus, ProposedLine
from motor.pgr_docx_utils import unique_cells
from motor.write_pgr import apply_lines_to_pgr, _find_insert_anchor, _row_category

AMENDO = Path(__file__).resolve().parents[2] / "modelos" / "PGR-Amendo.docx"


def _row_fill_cell(cell) -> str:
    shd = cell._tc.tcPr.find(qn("w:shd")) if cell._tc.tcPr is not None else None
    return shd.get(qn("w:fill"), "") if shd is not None else ""


def _line(*, ghe: str, table: int, psico_idx: int | None = None) -> ProposedLine:
    return ProposedLine(
        ghe_numero=ghe,
        ghe_nome=f"GHE {ghe}",
        setor_pgr="Teste",
        funcoes=[],
        categoria="Ergonômico (Psicossocial)",
        agente="Demandas Quantitativas Elevadas",
        exposicao="Habitual",
        causa_fonte="Organização do trabalho",
        trajetoria="Demanda",
        danos="Fadiga mental",
        grau_exposicao=3,
        grau_efeito=3,
        potencial="Moderado",
        controles="Revisar metas",
        evidencias=[],
        status=LineStatus.PROPOSTA,
        hazard_id="demandas_pressao_temporal",
        match_score=1.0,
        matched_from="test",
        n_respondentes=5,
        action="insert_after_psico",
        aprho_table_index=table,
        psico_row_index=psico_idx,
    )


def _row_fill(row) -> str:
    return _row_fill_cell(row.cells[0])


def test_insert_anchor_before_acidentes_amendo():
    if not AMENDO.exists():
        return

    doc = Document(str(AMENDO))
    table = doc.tables[10]
    anchor = _find_insert_anchor(table)
    cat = _row_category(table.rows[anchor])
    assert "psicos" in cat.lower() or "ergon" in cat.lower()
    assert "mec" not in cat.lower()
    assert "acident" not in cat.lower()


def test_insert_new_line_yellow_after_psico():
    if not AMENDO.exists():
        return

    lines = [_line(ghe="99", table=10, psico_idx=None)]
    with tempfile.TemporaryDirectory() as tmp:
        out = Path(tmp) / "out.docx"
        apply_lines_to_pgr(AMENDO, out, lines, only_accepted=False)
        doc = Document(str(out))
        table = doc.tables[10]
        psico_rows = [
            i
            for i, row in enumerate(table.rows)
            if "psicos" in _row_category(row).lower()
        ]
        assert len(psico_rows) >= 2
        new_idx = psico_rows[-1]
        assert _row_fill(table.rows[new_idx]).upper() == "FFFF00"
        # demais colunas sem cor de fundo
        uniq = unique_cells(table.rows[new_idx])
        for i, (_, cell) in enumerate(uniq):
            if i not in (0, 8):
                assert _row_fill_cell(cell) in ("", "auto", "FFFFFF")


if __name__ == "__main__":
    test_insert_anchor_before_acidentes_amendo()
    test_insert_new_line_yellow_after_psico()
    print("ok")
