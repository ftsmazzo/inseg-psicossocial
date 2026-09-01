"""Testes de preservação de formatação nas narrativas psicossociais."""

from __future__ import annotations

import tempfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from motor.pgr_narrative import apply_psicossocial_narratives

MAESTRALLE = Path(__file__).resolve().parents[2] / "modelos" / "PGR-Maestralle.docx"
AMENDO = Path(__file__).resolve().parents[2] / "modelos" / "PGR-Amendo.docx"

_MAX_SPACE_AFTER_BODY = 76200  # twips — limite em seções descritivas (não referências)


def _sect_count(doc: Document) -> int:
    return sum(
        1
        for p in doc.paragraphs
        if p._p.pPr is not None and p._p.pPr.sectPr is not None
    )


def _run_fonts(paragraph) -> set[str]:
    out: set[str] = set()
    for run in paragraph.runs:
        if not run.text.strip():
            continue
        r_pr = run._element.find(qn("w:rPr"))
        if r_pr is None:
            out.add("inherit")
            continue
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            out.add("inherit")
        else:
            out.add(r_fonts.get(qn("w:ascii")) or r_fonts.get(qn("w:hAnsi")) or "inherit")
    return out


def _find_para(doc: Document, needle: str):
    for i, p in enumerate(doc.paragraphs):
        if needle.lower() in p.text.lower():
            return i, p
    return None, None


def _assert_inserted_no_sect_or_inflated(doc: Document, needles: tuple[str, ...]) -> None:
    for needle in needles:
        _, p = _find_para(doc, needle)
        assert p is not None, f"parágrafo não encontrado: {needle}"
        if p._p.pPr is not None and p._p.pPr.sectPr is not None:
            raise AssertionError(f"sectPr em inserido: {needle}")
        sa = p.paragraph_format.space_after
        if sa is not None and sa > _MAX_SPACE_AFTER_BODY:
            raise AssertionError(f"space_after inflado ({sa}) em: {needle}")


def _assert_no_new_sect_breaks(before: Document, after: Document) -> None:
    assert _sect_count(after) == _sect_count(before)


def _can_open_docx(path: Path) -> bool:
    if not path.exists():
        return False
    try:
        Document(str(path))
        return True
    except Exception:
        return False


def test_narratives_preserve_structure_maestralle():
    if not _can_open_docx(MAESTRALLE):
        return

    before = Document(str(MAESTRALLE))
    doc = Document(str(MAESTRALLE))
    apply_psicossocial_narratives(doc)

    with tempfile.TemporaryDirectory() as tmp:
        out = Path(tmp) / "out.docx"
        doc.save(str(out))
        after = Document(str(out))

    _assert_no_new_sect_breaks(before, after)
    _assert_inserted_no_sect_or_inflated(
        after,
        (
            "Para os fatores psicossociais, a avaliação foi realizada",
            "O acompanhamento da saúde ocupacional deverá considerar",
            "Em casos de crise emocional, mal súbito",
            "Manter a calma e afastar o trabalhador",
        ),
    )


def test_inserted_body_text_no_explicit_cambria_maestralle():
    if not _can_open_docx(MAESTRALLE):
        return

    doc = Document(str(MAESTRALLE))
    apply_psicossocial_narratives(doc)

    _, p55 = _find_para(doc, "Para os fatores psicossociais, a avaliação foi realizada")
    assert p55 is not None
    assert p55.style.name == "Body Text"
    assert "Cambria" not in _run_fonts(p55)

    _, p53 = _find_para(
        doc,
        "O acompanhamento da saúde ocupacional deverá considerar também os fatores psicossociais",
    )
    assert p53 is not None
    assert p53.style.name == "List Paragraph"
    assert "Cambria" not in _run_fonts(p53)


def test_plano_emergencia_intro_body_style_maestralle():
    if not _can_open_docx(MAESTRALLE):
        return

    doc = Document(str(MAESTRALLE))
    apply_psicossocial_narratives(doc)

    _, intro = _find_para(doc, "Em casos de crise emocional, mal súbito")
    assert intro is not None
    assert intro.style.name == "Body Text"
    assert not (intro._p.pPr is not None and intro._p.pPr.numPr is not None)

    _, bullet = _find_para(doc, "Manter a calma e afastar o trabalhador")
    assert bullet is not None
    assert bullet._p.pPr is not None and bullet._p.pPr.numPr is not None


def test_references_no_cambria_explicit_maestralle():
    if not _can_open_docx(MAESTRALLE):
        return

    doc = Document(str(MAESTRALLE))
    apply_psicossocial_narratives(doc)

    _, ref = _find_para(doc, "Relatório de Avaliação Psicossocial (SSOS)")
    assert ref is not None
    assert "Cambria" not in _run_fonts(ref)


def test_append_5_2_no_new_paragraph_maestralle():
    if not _can_open_docx(MAESTRALLE):
        return

    before = Document(str(MAESTRALLE))
    count_before = len(before.paragraphs)

    doc = Document(str(MAESTRALLE))
    apply_psicossocial_narratives(doc)

    # append 5.2 não cria parágrafo; inserts criam 10 novos (1+1+1+6+2)
    assert len(doc.paragraphs) == count_before + 10


def test_narratives_amendo_if_present():
    if not _can_open_docx(AMENDO):
        return

    before = Document(str(AMENDO))
    doc = Document(str(AMENDO))
    apply_psicossocial_narratives(doc)
    _assert_no_new_sect_breaks(before, doc)


if __name__ == "__main__":
    test_narratives_preserve_structure_maestralle()
    test_inserted_body_text_no_explicit_cambria_maestralle()
    test_plano_emergencia_intro_body_style_maestralle()
    test_references_no_cambria_explicit_maestralle()
    test_append_5_2_no_new_paragraph_maestralle()
    test_narratives_amendo_if_present()
    print("ok")
