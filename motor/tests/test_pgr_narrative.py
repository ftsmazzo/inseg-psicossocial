"""Testes dos trechos narrativos psicossociais no PGR."""

from __future__ import annotations

import tempfile
from pathlib import Path

from docx import Document

from motor.pgr_narrative import apply_psicossocial_narratives

MODEL = Path(__file__).resolve().parents[2] / "modelos" / "PGR-Maestralle.docx"


def test_apply_narratives_maestralle_idempotent():
    if not MODEL.exists():
        return

    with tempfile.TemporaryDirectory() as tmp:
        out = Path(tmp) / "out.docx"
        doc = Document(str(MODEL))
        r1 = apply_psicossocial_narratives(doc)
        assert r1["5.2"] == "applied"
        assert r1["5.3"] == "applied"
        assert r1["5.5"] == "applied"
        assert r1["plano_emergencia_psico"] == "applied"
        assert r1["documentos_referencia_psico"] == "applied"
        doc.save(str(out))

        doc2 = Document(str(out))
        text = "\n".join(p.text for p in doc2.paragraphs)
        assert "fatores psicossociais relacionados à organização do trabalho" in text
        assert "ferramenta específica de levantamento e análise" in text
        assert "crise emocional, mal súbito" in text
        assert "Relatório de Avaliação Psicossocial (SSOS)" in text

        r2 = apply_psicossocial_narratives(doc2)
        assert all(v == "skipped" for v in r2.values())


if __name__ == "__main__":
    test_apply_narratives_maestralle_idempotent()
    print("ok")

