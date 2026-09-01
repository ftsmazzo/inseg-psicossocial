"""Fixture mínima com âncoras reais do PGR Inseg para testes de narrativa."""

from __future__ import annotations

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


def _add_numpr(p: Paragraph, num_id: int = 12) -> None:
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)


def _add_spacing_after(p: Paragraph, twips: int) -> None:
    p_pr = p._p.get_or_add_pPr()
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:after"), str(twips))
    p_pr.append(sp)


def build_minimal_narrative_doc() -> Document:
    """Documento sintético com âncoras e estilos típicos do PGR Inseg."""
    doc = Document()

    doc.add_paragraph(
        "A organização deve avaliar os riscos ocupacionais relativos aos perigos identificados.",
        style="Body Text",
    )

    p53 = doc.add_paragraph(
        "A organização deve desenvolver ações em saúde ocupacional dos trabalhadores.",
        style="List Paragraph",
    )
    _add_numpr(p53)
    _add_spacing_after(p53, 76200)
    next53 = doc.add_paragraph(
        "O controle da saúde dos empregados deve ser um processo preventivo.",
        style="List Paragraph",
    )
    _add_numpr(next53)

    doc.add_paragraph(
        "A organização deve selecionar as ferramentas e técnicas de avaliação de riscos "
        "adequadas ao risco ou circunstância.",
        style="Body Text",
    )
    _add_spacing_after(doc.paragraphs[-1], 635)

    anchor = doc.add_paragraph(
        "Em caso de acidente de trabalho de origem elétrica deverão ser seguidos os "
        "procedimentos especiais abaixo:",
        style="Heading 1",
    )
    _add_spacing_after(anchor, 76200)
    bullet = doc.add_paragraph(
        "Corte imediatamente a corrente elétrica, desligando a ficha do aparelho.",
        style="Heading 1",
    )
    _add_numpr(bullet)
    last_elec = doc.add_paragraph(
        "Se a vítima não der sinais de vida, depois de desligar a corrente elétrica faça RCP.",
        style="Heading 1",
    )
    _add_numpr(last_elec)
    _add_spacing_after(last_elec, 76200)
    doc.add_paragraph("EVACUAÇÃO:", style="Heading 1")

    doc.add_paragraph("DOCUMENTOS DE REFERÊNCIA", style="List Paragraph")
    nr9 = doc.add_paragraph(
        "Norma Regulamentadora nº 9 da Portaria 3.214/78, alterada pela Portaria SEPRT n.º 915/2019;",
        style="List Paragraph",
    )
    _add_numpr(nr9)
    run = nr9.runs[0]
    r_pr = run._element.get_or_add_rPr()
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Cambria")
    fonts.set(qn("w:hAnsi"), "Cambria")
    r_pr.append(fonts)
    _add_spacing_after(nr9, 152400)

    return doc
