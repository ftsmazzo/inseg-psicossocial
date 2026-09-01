"""Trechos descritivos psicossociais fixos nos PGR Inseg (mecânico, idempotente)."""

from __future__ import annotations

import logging
from dataclasses import dataclass
from typing import Literal

from docx import Document
from docx.text.paragraph import Paragraph

from motor.pgr_docx_utils import (
    append_to_paragraph,
    insert_paragraph_after,
    norm_text,
    paragraph_has_numpr,
)

logger = logging.getLogger(__name__)

Mode = Literal["append", "insert_after"]


def _norm(text: str) -> str:
    return norm_text(text)


def _doc_contains(doc: Document, marker: str) -> bool:
    m = _norm(marker)
    if not m:
        return False
    for para in doc.paragraphs:
        if m in _norm(para.text):
            return True
    return False


def _paragraph_index(doc: Document, para: Paragraph) -> int | None:
    for i, p in enumerate(doc.paragraphs):
        if p._p is para._p:
            return i
    return None


def _find_paragraph(doc: Document, *anchor_substrings: str) -> Paragraph | None:
    if not anchor_substrings:
        return None
    needles = [_norm(s) for s in anchor_substrings]
    for para in doc.paragraphs:
        hay = _norm(para.text)
        if all(n in hay for n in needles):
            return para
    return None


def _spacing_reference(doc: Document, after: Paragraph) -> Paragraph:
    """Parágrafo cujo espaçamento copiar — só vizinho da mesma seção/lista."""
    idx = _paragraph_index(doc, after)
    if idx is None:
        return after
    after_style = after.style.name if after.style else ""
    after_numpr = paragraph_has_numpr(after)
    if idx + 1 < len(doc.paragraphs):
        nxt = doc.paragraphs[idx + 1]
        if nxt.text.strip():
            nxt_style = nxt.style.name if nxt.style else ""
            nxt_numpr = paragraph_has_numpr(nxt)
            if nxt_style == after_style or (after_numpr and nxt_numpr):
                return nxt
    return after


def _find_neighbor_body_text(doc: Document, anchor_idx: int) -> Paragraph | None:
    """Parágrafo Body Text mais próximo — fallback para intro descritiva."""
    for i in range(anchor_idx, max(anchor_idx - 30, -1), -1):
        para = doc.paragraphs[i]
        if para.text.strip() and (para.style.name if para.style else "") == "Body Text":
            return para
    for i in range(anchor_idx, min(anchor_idx + 15, len(doc.paragraphs))):
        para = doc.paragraphs[i]
        if para.text.strip() and (para.style.name if para.style else "") == "Body Text":
            return para
    return None


def _find_last_references_anchor(doc: Document) -> Paragraph | None:
    last_heading: int | None = None
    for i, para in enumerate(doc.paragraphs):
        if _norm(para.text.replace("\t", "")) == _norm("DOCUMENTOS DE REFERÊNCIA"):
            last_heading = i
    if last_heading is None:
        return None

    insert_after = doc.paragraphs[last_heading]
    for i in range(last_heading + 1, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else ""
        if style.startswith("Heading") and _norm(text) != _norm("DOCUMENTOS DE REFERÊNCIA"):
            break
        insert_after = para
    return insert_after


def _find_evacuacao_paragraph(doc: Document) -> Paragraph | None:
    for para in doc.paragraphs:
        t = _norm(para.text)
        if t.startswith("evacua") and len(t) < 40:
            return para
    return None


def _last_electrical_bullet_before_evacuacao(doc: Document, anchor: Paragraph) -> Paragraph:
    """Último item numerado do bloco elétrico, imediatamente antes de EVACUAÇÃO."""
    idx = _paragraph_index(doc, anchor)
    if idx is None:
        return anchor

    evac_idx: int | None = None
    for i in range(idx + 1, len(doc.paragraphs)):
        t = _norm(doc.paragraphs[i].text)
        if t.startswith("evacua") and len(t) < 40:
            evac_idx = i
            break

    if evac_idx is not None and evac_idx > idx + 1:
        return doc.paragraphs[evac_idx - 1]

    insert_after = anchor
    for i in range(idx + 1, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        t = _norm(para.text)
        if not t:
            continue
        if t.startswith("evacua") or "incêndio" in t or "incendio" in t:
            break
        if "crise emocional" in t:
            break
        insert_after = para
    return insert_after


def _first_numbered_bullet_after(doc: Document, anchor: Paragraph) -> Paragraph:
    idx = _paragraph_index(doc, anchor)
    if idx is None:
        return anchor
    for i in range(idx + 1, min(idx + 8, len(doc.paragraphs))):
        para = doc.paragraphs[i]
        if paragraph_has_numpr(para) and len(para.text.strip()) > 15:
            return para
    return anchor


@dataclass(frozen=True)
class _Patch:
    patch_id: str
    mode: Mode
    anchors: tuple[str, ...]
    marker: str
    append_suffix: str = ""
    insert_lines: tuple[str, ...] = ()
    insert_style_from: str | None = None
    plano_emergencia_block: bool = False


_PATCHES: tuple[_Patch, ...] = (
    _Patch(
        patch_id="5.2",
        mode="append",
        anchors=("avaliar os riscos ocupacionais relativos aos perigos identificados",),
        marker="fatores psicossociais relacionados à organização do trabalho, conforme previsto na legislação vigente",
        append_suffix=(
            " e os fatores psicossociais relacionados à organização do trabalho, "
            "conforme previsto na legislação vigente e nas diretrizes de gestão de "
            "Segurança e Saúde no Trabalho adotadas pela organização."
        ),
    ),
    _Patch(
        patch_id="5.3",
        mode="insert_after",
        anchors=("desenvolver ações em saúde ocupacional dos trabalhadores",),
        marker="O acompanhamento da saúde ocupacional deverá considerar também os fatores psicossociais",
        insert_lines=(
            "O acompanhamento da saúde ocupacional deverá considerar também os fatores "
            "psicossociais relacionados ao trabalho, mediante monitoramento periódico e "
            "reavaliações sempre que ocorrerem alterações significativas na organização do trabalho.",
        ),
        insert_style_from="desenvolver ações em saúde ocupacional dos trabalhadores",
    ),
    _Patch(
        patch_id="5.5",
        mode="insert_after",
        anchors=("ferramentas e técnicas de avaliação de riscos", "adequadas ao risco ou circunstância"),
        marker="ferramenta específica de levantamento e análise dos fatores relacionados à organização do trabalho",
        insert_lines=(
            "Para os fatores psicossociais, a avaliação foi realizada por meio de ferramenta "
            "específica de levantamento e análise dos fatores relacionados à organização do "
            "trabalho, contemplando aspectos como demanda, controle, esforço e recompensa, "
            "sendo os resultados utilizados para subsidiar a identificação dos perigos, a "
            "avaliação dos riscos e a definição das medidas de prevenção aplicáveis.",
        ),
        insert_style_from="ferramentas e técnicas de avaliação de riscos",
    ),
    _Patch(
        patch_id="plano_emergencia_psico",
        mode="insert_after",
        anchors=("acidente de trabalho de origem elétrica", "procedimentos especiais"),
        marker="Em casos de crise emocional, mal súbito ou outras ocorrências relacionadas a fatores psicossociais",
        insert_lines=(
            "Em casos de crise emocional, mal súbito ou outras ocorrências relacionadas a "
            "fatores psicossociais que possam comprometer a segurança do trabalhador, "
            "deverão ser adotadas as seguintes medidas:",
            "Manter a calma e afastar o trabalhador de fontes de risco;",
            "Comunicar imediatamente o superior responsável;",
            "Encaminhar o trabalhador para local seguro e tranquilo;",
            "Acionar atendimento médico quando necessário;",
            "Registrar a ocorrência para acompanhamento e adoção das medidas cabíveis.",
        ),
        plano_emergencia_block=True,
    ),
    _Patch(
        patch_id="documentos_referencia_psico",
        mode="insert_after",
        anchors=(),
        marker="Relatório de Avaliação Psicossocial (SSOS)",
        insert_lines=(
            "Relatório de Avaliação Psicossocial (SSOS);",
            "Resultados consolidados da pesquisa de fatores psicossociais aplicada aos trabalhadores.",
        ),
        insert_style_from="Norma Regulamentadora nº 9",
    ),
)


def _insert_plano_emergencia_psico(doc: Document, anchor: Paragraph, lines: tuple[str, ...]) -> None:
    """Após acidentes elétricos e antes de EVACUAÇÃO — intro + marcadores numerados."""
    if not lines:
        return
    insert_after = _last_electrical_bullet_before_evacuacao(doc, anchor)
    bullet_tpl = _first_numbered_bullet_after(doc, anchor)
    anchor_idx = _paragraph_index(doc, anchor) or 0
    intro_tpl = _find_neighbor_body_text(doc, anchor_idx) or anchor

    intro, *bullets = list(lines)
    current = insert_paragraph_after(
        intro_tpl,
        intro,
        insert_after,
        keep_numpr=False,
        spacing_from=insert_after,
    )
    bullet_spacing = insert_after
    for item in bullets:
        current = insert_paragraph_after(
            bullet_tpl,
            item,
            current,
            keep_numpr=True,
            spacing_from=bullet_spacing,
        )
        bullet_spacing = current


def apply_psicossocial_narratives(doc: Document) -> dict[str, str | bool]:
    results: dict[str, str | bool] = {}

    for patch in _PATCHES:
        if _doc_contains(doc, patch.marker):
            results[patch.patch_id] = "skipped"
            continue

        if patch.patch_id == "documentos_referencia_psico":
            anchor = _find_last_references_anchor(doc)
            if anchor is None:
                results[patch.patch_id] = "anchor_not_found"
                logger.warning("PGR narrative: seção DOCUMENTOS DE REFERÊNCIA não encontrada")
                continue
            tpl = _find_paragraph(doc, patch.insert_style_from or "") or anchor
            current = anchor
            for line in patch.insert_lines:
                current = insert_paragraph_after(
                    tpl,
                    line,
                    current,
                    keep_numpr=paragraph_has_numpr(tpl),
                    spacing_from=_spacing_reference(doc, anchor),
                )
            results[patch.patch_id] = "applied"
            continue

        anchor = _find_paragraph(doc, *patch.anchors)
        if anchor is None:
            results[patch.patch_id] = "anchor_not_found"
            logger.warning(
                "PGR narrative %s: âncora não encontrada (%s)",
                patch.patch_id,
                patch.anchors,
            )
            continue

        if patch.mode == "append":
            append_to_paragraph(anchor, patch.append_suffix)
            results[patch.patch_id] = "applied"
            continue

        if patch.plano_emergencia_block:
            _insert_plano_emergencia_psico(doc, anchor, patch.insert_lines)
            results[patch.patch_id] = "applied"
            continue

        tpl = anchor
        if patch.insert_style_from:
            src = _find_paragraph(doc, patch.insert_style_from)
            if src is not None:
                tpl = src
        current = anchor
        for line in patch.insert_lines:
            current = insert_paragraph_after(
                tpl,
                line,
                current,
                keep_numpr=paragraph_has_numpr(tpl),
                spacing_from=_spacing_reference(doc, anchor),
            )
        results[patch.patch_id] = "applied"

    return results
