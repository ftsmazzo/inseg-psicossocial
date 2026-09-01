from __future__ import annotations

import re
from pathlib import Path

from docx import Document

from motor.models import AprhoCell, GheBlock, PgrModel
from motor.textutil import normalize, parse_float_br


def _cell_text(cell) -> str:
    return "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip()).strip()


def _row_cells(row) -> list[str]:
    return [_cell_text(c) for c in row.cells]


INSEG_MATRIZ_DEFAULT: dict[int, dict[int, str]] = {
    5: {1: "Baixo", 2: "Moderado", 3: "Alto", 4: "Muito Alto", 5: "Muito Alto"},
    4: {1: "Baixo", 2: "Moderado", 3: "Moderado", 4: "Alto", 5: "Muito Alto"},
    3: {1: "Baixo", 2: "Baixo", 3: "Moderado", 4: "Moderado", 5: "Alto"},
    2: {1: "Muito Baixo", 2: "Baixo", 3: "Baixo", 4: "Moderado", 5: "Moderado"},
    1: {1: "Muito Baixo", 2: "Muito Baixo", 3: "Baixo", 4: "Baixo", 5: "Baixo"},
}


def lookup_potencial(ge: int, ges: int, matriz: dict[int, dict[int, str]] | None = None) -> str:
    m = matriz or INSEG_MATRIZ_DEFAULT
    return m.get(ge, {}).get(ges, "Não determinado")


def parse_pgr_inseg(path: str | Path) -> PgrModel:
    path = Path(path)
    doc = Document(str(path))

    razao, cnpj = _extract_empresa(doc)
    matriz = _extract_matriz(doc) or dict(INSEG_MATRIZ_DEFAULT)
    ge_desc, ges_desc = _extract_escalas(doc)
    ghes = _extract_ghes(doc)

    return PgrModel(
        source_file=path.name,
        razao_social=razao,
        cnpj=cnpj,
        ghes=ghes,
        matriz=matriz,
        grau_exposicao_desc=ge_desc,
        grau_efeito_desc=ges_desc,
    )


def _extract_empresa(doc: Document) -> tuple[str, str]:
    razao, cnpj = "", ""
    for table in doc.tables[:5]:
        for row in table.rows:
            cells = _row_cells(row)
            if len(cells) >= 2:
                label = normalize(cells[0])
                if "razao social" in label or label == "razao social":
                    razao = cells[1].strip()
                if label == "cnpj":
                    cnpj = cells[1].strip()
    if not cnpj:
        blob = "\n".join(p.text for p in doc.paragraphs[:80])
        m = re.search(r"(\d{2}\.\d{3}\.\d{3}/\d{4}-\d{2})", blob)
        if m:
            cnpj = m.group(1)
    return razao, cnpj


def _extract_escalas(doc: Document) -> tuple[dict[int, str], dict[int, str]]:
    ge_desc: dict[int, str] = {}
    ges_desc: dict[int, str] = {}
    for table in doc.tables:
        rows = [_row_cells(r) for r in table.rows]
        if not rows:
            continue
        h = normalize(" ".join(rows[0]))
        if "caracteristicas da exposicao" in h or ("grau" in h and "exposicao" in h and len(rows[0]) <= 3):
            for r in rows[1:]:
                if not r:
                    continue
                m = re.match(r"(\d+)", r[0])
                if m:
                    ge_desc[int(m.group(1))] = (r[1] if len(r) > 1 else r[0])[:300]
        if "efeitos" in h and "saude" in h:
            for r in rows[1:]:
                if not r:
                    continue
                m = re.match(r"(\d+)", r[0])
                if m:
                    ges_desc[int(m.group(1))] = (r[1] if len(r) > 1 else r[0])[:300]
    return ge_desc, ges_desc


def _extract_matriz(doc: Document) -> dict[int, dict[int, str]] | None:
    for table in doc.tables:
        rows = [_row_cells(r) for r in table.rows]
        if len(rows) < 5:
            continue
        joined = normalize(" ".join(rows[0]))
        if "grau de exposicao" not in joined:
            continue
        # Expect rows with GE label and 5 potencial values
        matriz: dict[int, dict[int, str]] = {}
        for r in rows:
            if not r:
                continue
            m = re.search(r"\b([1-5])\b", r[0]) if r[0] else None
            # cells may be merged; find numeric ge in first cells
            ge = None
            vals: list[str] = []
            for c in r:
                cm = re.fullmatch(r"[1-5]", c.strip()) if c else None
                if cm and ge is None and "grau" not in normalize(c):
                    # could be ge column value sitting alone
                    pass
            # Pattern from analysis: first meaningful number after label is GE 5..1
            nums = re.findall(r"\b([1-5])\b", " ".join(r[:3]))
            potenciais = [
                c.strip()
                for c in r
                if c.strip()
                in {"Muito Baixo", "Baixo", "Moderado", "Alto", "Muito Alto"}
                or normalize(c) in {
                    "muito baixo",
                    "baixo",
                    "moderado",
                    "alto",
                    "muito alto",
                }
            ]
            if len(potenciais) == 5 and nums:
                ge = int(nums[0])
                matriz[ge] = {i + 1: potenciais[i] for i in range(5)}
        if len(matriz) >= 5:
            return matriz
    return None


def _extract_ghes(doc: Document) -> list[GheBlock]:
    # Map paragraph index of "Caracterização do GHE XX" to following characterization + APRHO tables
    headings: list[tuple[int, str, str]] = []
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        m = re.match(
            r"Caracteriza[cç][aã]o do GHE\s*(\d+)\s*[–\-—:]\s*(.+)",
            t,
            re.I,
        )
        if m:
            headings.append((i, m.group(1), m.group(2).strip()))

    # Build table inventory with types
    tables_meta: list[dict] = []
    for ti, table in enumerate(doc.tables):
        rows = [_row_cells(r) for r in table.rows]
        kind = "other"
        if rows:
            h0 = normalize(" ".join(rows[0][:4]))
            h1 = normalize(" ".join(rows[1][:4])) if len(rows) > 1 else ""
            if "setor" in h0 and ("cargo" in h0 or "funcao" in h0):
                kind = "caracterizacao"
            elif "categoria" in h1 and "agente" in h1:
                kind = "aprho"
            elif "categoria" in h0 and "agente" in h0:
                kind = "aprho"
        tables_meta.append({"index": ti, "kind": kind, "rows": rows})

    # Associate: after each heading, next caracterizacao then next aprho
    # Heuristic by document order: walk tables in order interleaved with headings via XML is hard;
    # Use sequential pairing: nth caracterizacao + nth aprho for nth GHE heading.
    caracts = [t for t in tables_meta if t["kind"] == "caracterizacao"]
    aprhos = [t for t in tables_meta if t["kind"] == "aprho"]

    ghes: list[GheBlock] = []
    n = min(len(headings), len(aprhos))
    for i in range(n):
        _pi, num, nome = headings[i]
        setor, funcoes, atividade, ambiente = "", [], "", ""
        if i < len(caracts):
            rows = caracts[i]["rows"]
            if len(rows) > 1:
                r = rows[1]
                setor = r[0] if len(r) > 0 else ""
                func_raw = r[1] if len(r) > 1 else nome
                funcoes = _split_funcoes(func_raw)
                atividade = r[2] if len(r) > 2 else ""
                ambiente = r[3] if len(r) > 3 else ""
        if not funcoes:
            funcoes = _split_funcoes(nome)

        aprho = aprhos[i]
        psico = _find_psico_row(aprho["rows"])
        cats = []
        for r in aprho["rows"][2:]:
            if r and r[0] and r[0] not in cats and len(r[0]) < 80:
                cats.append(r[0])

        ghes.append(
            GheBlock(
                numero=num.zfill(2) if num.isdigit() else num,
                nome=nome,
                setor=setor,
                funcoes=funcoes,
                atividade=atividade,
                ambiente=ambiente,
                aprho_table_index=aprho["index"],
                psico_row=psico,
                all_categories=cats,
            )
        )
    return ghes


def _split_funcoes(text: str) -> list[str]:
    parts = re.split(r"[/|;]", text)
    return [p.strip() for p in parts if p.strip()]


def _find_psico_row(rows: list[list[str]]) -> AprhoCell | None:
    for idx, r in enumerate(rows):
        if len(r) < 9:
            continue
        cat = r[0]
        if "psicos" in normalize(cat):
            return AprhoCell(
                categoria=r[0],
                agente=r[1] if len(r) > 1 else "",
                exposicao=r[2] if len(r) > 2 else "",
                causa_fonte=r[3] if len(r) > 3 else "",
                trajetoria=r[4] if len(r) > 4 else "",
                danos=r[5] if len(r) > 5 else "",
                grau_exposicao=r[6] if len(r) > 6 else "",
                grau_efeito=r[7] if len(r) > 7 else "",
                potencial=r[8] if len(r) > 8 else "",
                controles=r[9] if len(r) > 9 else "",
                row_index=idx,
            )
    return None
