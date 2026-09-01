"""Utilitários compartilhados para formatação mecânica de PGR DOCX."""

from __future__ import annotations

import copy
import re
import unicodedata

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


def norm_text(text: str) -> str:
    t = unicodedata.normalize("NFKC", text or "")
    return re.sub(r"\s+", " ", t.replace("\xa0", " ")).strip().lower()


def set_cell_fill(cell, fill_hex: str | None) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if fill_hex:
        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), fill_hex.upper())
    elif shd is not None:
        tc_pr.remove(shd)


def clear_cell_fill(cell) -> None:
    set_cell_fill(cell, None)


def unique_cells(row) -> list[tuple[int, object]]:
    seen: set[int] = set()
    out: list[tuple[int, object]] = []
    for idx, cell in enumerate(row.cells):
        tc_id = id(cell._tc)
        if tc_id in seen:
            continue
        seen.add(tc_id)
        out.append((idx, cell))
    return out


def _run_has_drawing(run) -> bool:
    return bool(run._element.findall(f".//{qn('w:drawing')}"))


def _set_run_text_preserve_element(run, text: str) -> None:
    """Atualiza w:t sem recriar o run (python-docx invalida referências ao usar .text)."""
    for t_el in run._element.findall(qn("w:t")):
        t_el.text = text
        return
    run.text = text


def set_cell_text(cell, text: str) -> None:
    """Substitui texto visível preservando desenhos/watermark no parágrafo."""
    text = text or ""
    if not cell.paragraphs:
        cell.text = text
        return
    p0 = cell.paragraphs[0]
    runs = p0.runs
    if not runs:
        p0.text = text
        return

    text_idx: int | None = None
    empty_idx: int | None = None
    for i, run in enumerate(runs):
        if _run_has_drawing(run):
            continue
        if run.text.strip():
            text_idx = i
            break
        if empty_idx is None:
            empty_idx = i

    if text_idx is None:
        text_idx = empty_idx if empty_idx is not None else 0

    _set_run_text_preserve_element(runs[text_idx], text)

    for i, run in enumerate(p0.runs):
        if i == text_idx or _run_has_drawing(run):
            continue
        _set_run_text_preserve_element(run, "")

    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)


def _primary_text_run(paragraph: Paragraph):
    chosen = None
    for run in paragraph.runs:
        if _run_has_drawing(run):
            continue
        if run.text.strip():
            return run
        if chosen is None:
            chosen = run
    return chosen


def _set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    text_run = _primary_text_run(paragraph)
    if text_run is not None:
        _set_run_text_preserve_element(text_run, text)
        for run in paragraph.runs:
            if run._element is not text_run._element and not _run_has_drawing(run):
                _set_run_text_preserve_element(run, "")
    else:
        paragraph.text = text


def _clear_explicit_run_fonts(paragraph: Paragraph) -> None:
    """Remove rFonts explícito — herda Calibri (ou fonte do estilo) do tema."""
    for run in paragraph.runs:
        if _run_has_drawing(run):
            continue
        r_pr = run._element.find(qn("w:rPr"))
        if r_pr is None:
            continue
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is not None:
            r_pr.remove(r_fonts)


def _sanitize_paragraph_properties(
    p_el,
    *,
    keep_numpr: bool,
    spacing_from: Paragraph | None,
) -> None:
    """Remove sectPr/pageBreak e normaliza espaçamento — evita saltos e sobreposição."""
    p_pr = p_el.find(qn("w:pPr"))
    if p_pr is None:
        return

    sect = p_pr.find(qn("w:sectPr"))
    if sect is not None:
        p_pr.remove(sect)

    for tag in ("w:pageBreakBefore", "w:framePr", "w:keepNext", "w:keepLines"):
        el = p_pr.find(qn(tag))
        if el is not None:
            p_pr.remove(el)

    if not keep_numpr:
        numpr = p_pr.find(qn("w:numPr"))
        if numpr is not None:
            p_pr.remove(numpr)

    for tag in ("w:spacing", "w:ind"):
        el = p_pr.find(qn(tag))
        if el is not None:
            p_pr.remove(el)

    if spacing_from is not None:
        src_pr = spacing_from._p.find(qn("w:pPr"))
        if src_pr is not None:
            for tag in ("w:spacing", "w:ind"):
                src_el = src_pr.find(qn(tag))
                if src_el is not None:
                    p_pr.append(copy.deepcopy(src_el))


def append_to_paragraph(paragraph: Paragraph, suffix: str) -> None:
    """Acrescenta texto ao parágrafo existente sem criar parágrafo novo."""
    base = paragraph.text.rstrip()
    if base.endswith("."):
        base = base[:-1].rstrip()
    _set_paragraph_text(paragraph, f"{base}{suffix}")


def insert_paragraph_after(
    template: Paragraph,
    text: str,
    after: Paragraph,
    *,
    keep_numpr: bool = True,
    spacing_from: Paragraph | None = None,
    strip_explicit_fonts: bool = True,
) -> Paragraph:
    """
    Insere parágrafo clonando estilo/lista do template, com mínima interferência:
    sem sectPr, sem espaçamento inflado do template, fonte herdada do estilo.
    """
    spacing_from = after if spacing_from is None else spacing_from
    new_el = copy.deepcopy(template._p)
    _sanitize_paragraph_properties(
        new_el,
        keep_numpr=keep_numpr,
        spacing_from=spacing_from,
    )
    after._p.addnext(new_el)
    new_para = Paragraph(new_el, after._parent)
    _set_paragraph_text(new_para, text)
    if strip_explicit_fonts:
        _clear_explicit_run_fonts(new_para)
    return new_para


def clone_paragraph_after(template: Paragraph, text: str, after: Paragraph) -> Paragraph:
    """Compat — delega para insert_paragraph_after sanitizado."""
    return insert_paragraph_after(
        template,
        text,
        after,
        keep_numpr=paragraph_has_numpr(template),
        spacing_from=after,
        strip_explicit_fonts=True,
    )


def paragraph_has_numpr(paragraph: Paragraph) -> bool:
    p_pr = paragraph._p.pPr
    return p_pr is not None and p_pr.numPr is not None


_POTENCIAL_FILL = {
    "muito baixo": "92D050",
    "baixo": "92D050",
    "moderado": "FFFF00",
    "médio": "FFFF00",
    "medio": "FFFF00",
    "alto": "FFC000",
    "crítico": "FF0000",
    "critico": "FF0000",
}


def style_psico_aprho_row(row, *, potencial: str, categoria: str = "") -> None:
    """Só categoria (amarelo) e potencial (cor do risco); demais células sem fill."""
    pot_fill = _POTENCIAL_FILL.get(norm_text(potencial), "FFFF00")
    uniq = unique_cells(row)
    for i, (_, cell) in enumerate(uniq):
        if i == 0:
            set_cell_fill(cell, "FFFF00")
        elif i == 8:
            set_cell_fill(cell, pot_fill)
        else:
            clear_cell_fill(cell)
