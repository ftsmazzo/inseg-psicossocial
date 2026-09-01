from pathlib import Path
from docx import Document
import json
import re

BASE = Path(r"c:\Users\anjo_\OneDrive\Projetos-FabriaIA\psicossocial\modelos")
OUT = Path(r"c:\Users\anjo_\OneDrive\Projetos-FabriaIA\psicossocial\scripts\analysis_out")
OUT.mkdir(exist_ok=True)


def cell_text(cell: object) -> str:
    return "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip()).strip()


def dedupe_row(cells: list[str]) -> list[str]:
    out: list[str] = []
    for c in cells:
        if not out or out[-1] != c:
            out.append(c)
    return out


KEYWORDS = (
    "psicos",
    "ergon",
    "probab",
    "sever",
    "ghe",
    "agente",
    "invent",
    "exposição",
    "exposicao",
    "causa",
    "fonte",
    "trajet",
    "dano",
    "matriz",
    "grau",
)


def analyze_docx(path: Path) -> dict:
    doc = Document(str(path))
    paragraphs_sample: list[str] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t and len(paragraphs_sample) < 60:
            paragraphs_sample.append(t[:240])

    # Find headings that look like GHE / setor / inventário
    headings: list[str] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        low = t.lower()
        if any(
            k in low
            for k in (
                "ghe",
                "grupo homog",
                "inventário",
                "inventario",
                "matriz",
                "psicos",
                "ergonôm",
                "ergonom",
                "plano de ação",
                "plano de acao",
                "apr",
            )
        ):
            headings.append(t[:200])
            if len(headings) >= 80:
                break

    tables = []
    for ti, table in enumerate(doc.tables):
        rows_data = []
        for row in table.rows:
            cells = [cell_text(c) for c in row.cells]
            rows_data.append(dedupe_row(cells))

        flat = " || ".join(" | ".join(r) for r in rows_data[:8]).lower()
        interesting = any(k in flat for k in KEYWORDS) or any(
            any(k in c.lower() for k in KEYWORDS) for r in rows_data[:3] for c in r
        )

        tables.append(
            {
                "index": ti,
                "n_rows": len(table.rows),
                "n_cols": len(table.columns),
                "header": [c[:120] for c in (rows_data[0] if rows_data else [])],
                "row1": [c[:120] for c in (rows_data[1] if len(rows_data) > 1 else [])],
                "row2": [c[:120] for c in (rows_data[2] if len(rows_data) > 2 else [])],
                "row_last": [c[:120] for c in (rows_data[-1] if rows_data else [])],
                "interesting": interesting,
                "has_ergonomico": "ergon" in flat,
                "has_psicos": "psicos" in flat,
            }
        )

    return {
        "file": path.name,
        "n_tables": len(doc.tables),
        "paragraphs_sample": paragraphs_sample,
        "headings_hits": headings,
        "tables": tables,
    }


def dump_table_full(path: Path, table_index: int, max_rows: int = 30) -> list[list[str]]:
    doc = Document(str(path))
    table = doc.tables[table_index]
    rows = []
    for i, row in enumerate(table.rows):
        if i >= max_rows:
            break
        rows.append([cell_text(c)[:200] for c in row.cells])
    return rows


def main() -> None:
    # Start with Amendo (smaller)
    pgr_path = BASE / "PGR-Amendo.docx"
    print(f"Analyzing {pgr_path.name}...")
    info = analyze_docx(pgr_path)
    (OUT / "pgr_amendo_summary.json").write_text(
        json.dumps(info, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    print(f"n_tables={info['n_tables']}")
    print(f"heading hits={len(info['headings_hits'])}")
    for h in info["headings_hits"][:40]:
        print(f"  H: {h}")

    print("\nInteresting / large tables:")
    for t in info["tables"]:
        if t["interesting"] or t["n_rows"] >= 4:
            print(
                f"  T{t['index']}: {t['n_rows']}x{t['n_cols']} "
                f"ergo={t['has_ergonomico']} psico={t['has_psicos']}"
            )
            print(f"    header: {t['header'][:15]}")
            print(f"    row1: {t['row1'][:10]}")

    # Dump the most relevant inventory-like tables fully
    candidates = [
        t["index"]
        for t in info["tables"]
        if t["interesting"] or (t["n_rows"] >= 5 and t["n_cols"] >= 5)
    ]
    for idx in candidates[:12]:
        rows = dump_table_full(pgr_path, idx, max_rows=25)
        (OUT / f"pgr_amendo_table_{idx}.json").write_text(
            json.dumps(rows, ensure_ascii=False, indent=2), encoding="utf-8"
        )
        print(f"dumped table {idx} ({len(rows)} rows)")


if __name__ == "__main__":
    main()
