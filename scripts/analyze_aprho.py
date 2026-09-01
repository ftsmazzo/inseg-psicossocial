from pathlib import Path
from docx import Document
import json
import re

BASE = Path(r"c:\Users\anjo_\OneDrive\Projetos-FabriaIA\psicossocial\modelos")
OUT = Path(r"c:\Users\anjo_\OneDrive\Projetos-FabriaIA\psicossocial\scripts\analysis_out")


def cell_text(cell) -> str:
    return "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip()).strip()


def analyze_aprho(path: Path) -> dict:
    doc = Document(str(path))
    ghes = []
    for p in doc.paragraphs:
        t = p.text.strip()
        m = re.match(r"Caracteriza[cç][aã]o do GHE\s*(\d+)\s*[–\-—]\s*(.+)", t, re.I)
        if m:
            ghes.append({"num": m.group(1), "nome": m.group(2).strip()})

    # Find APRHO tables (10 cols, header with Categoria/Agente)
    aprhos = []
    for ti, table in enumerate(doc.tables):
        if len(table.columns) < 8:
            continue
        rows = []
        for row in table.rows:
            rows.append([cell_text(c) for c in row.cells])
        if len(rows) < 2:
            continue
        header_join = " ".join(rows[1] if len(rows) > 1 else rows[0]).lower()
        if "categoria" in header_join and "agente" in header_join:
            # extract unique categories and last ergonomic-like rows
            categories = []
            data_rows = []
            for r in rows[2:]:
                # first col often category (with merges, may repeat)
                cat = r[0] if r else ""
                agent = r[1] if len(r) > 1 else ""
                if cat and cat not in categories:
                    # only unique short cats
                    if len(cat) < 80:
                        categories.append(cat)
                data_rows.append(
                    {
                        "categoria": cat[:80],
                        "agente": agent[:120],
                        "exposicao": (r[2][:100] if len(r) > 2 else ""),
                        "grau_exp": (r[6][:40] if len(r) > 6 else ""),
                        "grau_efeito": (r[7][:40] if len(r) > 7 else ""),
                        "potencial": (r[8][:40] if len(r) > 8 else ""),
                        "controles": (r[9][:120] if len(r) > 9 else ""),
                    }
                )
            aprhos.append(
                {
                    "table_index": ti,
                    "n_rows": len(rows),
                    "header_row0": [c[:60] for c in rows[0][:10]],
                    "header_row1": [c[:60] for c in rows[1][:10]],
                    "categories": categories,
                    "data_rows": data_rows,
                }
            )

    # Matrix tables T3 T4 T5 T6 style
    matrix = {"grau_exposicao": [], "grau_efeito": [], "matriz": [], "acoes": []}
    for ti, table in enumerate(doc.tables):
        rows = [[cell_text(c) for c in row.cells] for row in table.rows]
        if not rows:
            continue
        h0 = " ".join(rows[0]).upper()
        if "CARACTER" in h0 and "EXPOSI" in h0:
            matrix["grau_exposicao"] = rows
        elif "EFEITOS" in h0 or "EFEITO" in h0:
            matrix["grau_efeito"] = rows
        elif "GRAU DE EXPOSI" in h0:
            matrix["matriz"] = [[c[:30] for c in r[:8]] for r in rows]
        elif "A" in h0 and "RECOMEND" in h0:
            matrix["acoes"] = [[c[:80] for c in r[:3]] for r in rows]

    return {
        "file": path.name,
        "ghes": ghes,
        "n_aprho": len(aprhos),
        "aprho_summaries": [
            {
                "table_index": a["table_index"],
                "n_rows": a["n_rows"],
                "categories": a["categories"],
                "n_data_rows": len(a["data_rows"]),
                "agents": [d["agente"] for d in a["data_rows"]],
                "last_3": a["data_rows"][-3:],
            }
            for a in aprhos
        ],
        "matrix": matrix,
        "full_first_aprho": aprhos[0] if aprhos else None,
    }


def main():
    for name in ["PGR-Amendo.docx", "PGR-Piuka.docx", "PGR-Polimetal.docx"]:
        path = BASE / name
        print(f"\n==== {name} ====")
        info = analyze_aprho(path)
        out = OUT / f"{path.stem}_aprho.json"
        out.write_text(json.dumps(info, ensure_ascii=False, indent=2), encoding="utf-8")
        print(f"GHEs ({len(info['ghes'])}):")
        for g in info["ghes"]:
            print(f"  {g['num']}: {g['nome']}")
        print(f"APRHO tables: {info['n_aprho']}")
        for a in info["aprho_summaries"]:
            print(
                f"  T{a['table_index']}: cats={a['categories']} agents={a['agents']}"
            )
        print("Matrix GE rows:", len(info["matrix"]["grau_exposicao"]))
        if info["matrix"]["grau_exposicao"]:
            for r in info["matrix"]["grau_exposicao"]:
                print(" ", r[:2])
        if info["matrix"]["grau_efeito"]:
            print("Matrix GES:")
            for r in info["matrix"]["grau_efeito"]:
                print(" ", r[:2])
        if info["matrix"]["matriz"]:
            print("Matriz:")
            for r in info["matrix"]["matriz"]:
                print(" ", r)


if __name__ == "__main__":
    main()
